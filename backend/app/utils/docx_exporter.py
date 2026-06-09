import os
import io
from datetime import datetime
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from app.models.cruise import CruisePlan

ROLE_LABELS = {
    "investigador_principal": "Investigador Principal",
    "coinvestigador": "Co-investigador",
    "tecnico": "Técnico",
    "estudiante": "Estudiante",
    "capitan": "Capitán",
    "primer_oficial": "Primer Oficial",
    "marinero": "Marinero",
    "jefe_maquinas": "Jefe de Máquinas",
    "medico": "Médico",
    "otro": "Otro"
}

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def style_table(table):
    # Set thin gray borders and alignment
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            '<w:tblBorders %s>'
            '  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
            '  <w:left w:val="none"/>'
            '  <w:right w:val="none"/>'
            '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
            '  <w:insideV w:val="none"/>'
            '</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders)

def format_date(d):
    if not d:
        return "—"
    # handles date or datetime
    if isinstance(d, datetime):
        return d.strftime("%d/%m/%Y a las %H:%M")
    return d.strftime("%d/%m/%Y")

def get_static_map_image(waypoints, departure_port=None, return_port=None):
    # Construct complete path points list
    points = []
    if departure_port and departure_port.latitude is not None and departure_port.longitude is not None:
        points.append(departure_port)
    points.extend(waypoints)
    if return_port and return_port.latitude is not None and return_port.longitude is not None:
        points.append(return_port)

    if not points:
        return None
    try:
        import urllib.request
        # 1. Polyline
        pl_coords = ",".join(f"{w.longitude:.5f},{w.latitude:.5f}" for w in points)
        pl_param = f"c:0A2647FF,w:3,{pl_coords}"
        
        # 2. Markers
        markers = []
        for i, w in enumerate(points):
            is_start = i == 0
            is_end = i == len(points) - 1 and len(points) > 1
            is_science = len(w.samples) > 0 if hasattr(w, 'samples') and w.samples else False
            
            if is_start:
                color = "gn" # green
            elif is_end:
                color = "rd" # red
            elif is_science:
                color = "or" # orange
            else:
                color = "bl" # blue
                
            markers.append(f"{w.longitude:.5f},{w.latitude:.5f},pm2{color}m{i+1}")
            
        pt_param = "~".join(markers)
        url = f"https://static-maps.yandex.ru/1.x/?l=map&size=600,380&pl={pl_param}&pt={pt_param}&lang=es_MX"
        
        headers = {"User-Agent": "Mozilla/5.0"}
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=5) as response:
            return io.BytesIO(response.read())
    except Exception as e:
        print(f"Error downloading static map: {e}")
        return None

def generate_docx(cruise: CruisePlan) -> io.BytesIO:
    doc = docx.Document()
    
    # Page setup (margins)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base styling (font)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 1. Header (Three column table with logos on the corners and text in the middle)
    app_dir = os.path.dirname(os.path.dirname(__file__))
    logo_cicese = os.path.join(app_dir, "cicese_logo.jpg")
    logo_siae = os.path.join(app_dir, "SIAE_Logo_shield_Isotipo_light_512x512.png")

    t_header = doc.add_table(rows=1, cols=3)
    t_header.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Borderless except bottom line
    tblPr = t_header._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            '<w:tblBorders %s>'
            '  <w:top w:val="none"/>'
            '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>'
            '  <w:left w:val="none"/>'
            '  <w:right w:val="none"/>'
            '  <w:insideH w:val="none"/>'
            '  <w:insideV w:val="none"/>'
            '</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders)

    cell_left, cell_mid, cell_right = t_header.rows[0].cells
    cell_left.width = Inches(2.0)
    cell_mid.width = Inches(3.0)
    cell_right.width = Inches(1.5)

    # Margins and vertical centering alignment
    for cell in (cell_left, cell_mid, cell_right):
        set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
        tcPr = cell._element.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

    # Left: CICESE Logo
    if os.path.exists(logo_cicese):
        try:
            p_left = cell_left.paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left.paragraph_format.space_after = Pt(0)
            p_left.paragraph_format.space_before = Pt(0)
            r_left = p_left.add_run()
            r_left.add_picture(logo_cicese, width=Inches(1.8))
        except Exception:
            pass

    # Middle: Text Info
    p_mid = cell_mid.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mid.paragraph_format.space_before = Pt(0)
    p_mid.paragraph_format.space_after = Pt(0)
    r_mid_1 = p_mid.add_run("SIAE — SISTEMA INTEGRAL DE\nADMINISTRACIÓN DE EMBARCACIONES\n")
    r_mid_1.bold = True
    r_mid_1.font.size = Pt(8)
    r_mid_1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    r_mid_2 = p_mid.add_run("CICESE — DEPARTAMENTO DE EMBARCACIONES OCEANOGRÁFICAS")
    r_mid_2.font.size = Pt(6.5)
    r_mid_2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Right: SIAE Logo (right aligned)
    if os.path.exists(logo_siae):
        try:
            p_right = cell_right.paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.paragraph_format.space_after = Pt(0)
            p_right.paragraph_format.space_before = Pt(0)
            r_right = p_right.add_run()
            r_right.add_picture(logo_siae, width=Inches(1.4))
        except Exception:
            pass

    # Space after header table
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)
    p_space.paragraph_format.space_after = Pt(0)
    p_space.paragraph_format.line_spacing = Pt(1)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PLAN DE CAMPAÑA DE INVESTIGACIÓN MARINA")
    run_title.bold = True
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = RGBColor(0x0A, 0x26, 0x47) # Navy blue

    # Folio / Cruise number
    p_folio = doc.add_paragraph()
    p_folio.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_folio = p_folio.add_run(f"Folio / Número de Crucero: {cruise.cruise_number or '—'}\n")
    run_folio.bold = True
    run_folio.font.size = Pt(12)
    run_folio.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # 2. General Data Table
    p_sec1 = doc.add_paragraph()
    run_sec1 = p_sec1.add_run("1. DATOS GENERALES DEL CRUCERO")
    run_sec1.bold = True
    run_sec1.font.size = Pt(12)
    run_sec1.font.color.rgb = RGBColor(0x0A, 0x26, 0x47)

    table_data = [
        ("Nombre del Crucero", cruise.name),
        ("Embarcación", cruise.vessel.name if cruise.vessel else "—"),
        ("Capitán", cruise.captain.full_name if cruise.captain else "Sin asignar"),
        ("Proyecto Asociado", cruise.project_name or "—"),
        ("Fuente de Financiamiento", cruise.funding_source or "—"),
        ("Área de Estudio", cruise.study_area or "—"),
        ("Disciplinas Científicas", cruise.disciplines or "—"),
        ("Puerto de Salida", cruise.departure_port or "—"),
        ("Fecha de Salida", format_date(cruise.departure_date)),
        ("Puerto de Regreso", cruise.return_port or "—"),
        ("Fecha de Regreso", format_date(cruise.return_date)),
        ("Millas Náuticas Planificadas", f"{cruise.planned_nm} mn" if cruise.planned_nm is not None else "—"),
    ]

    t_gen = doc.add_table(rows=0, cols=2)
    style_table(t_gen)

    for label, val in table_data:
        row = t_gen.add_row()
        cell_lbl, cell_val = row.cells
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        # Label cell
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(2)
        p_lbl.paragraph_format.space_before = Pt(2)
        r_lbl = p_lbl.add_run(label)
        r_lbl.bold = True
        r_lbl.font.size = Pt(10)
        set_cell_background(cell_lbl, "F0F4F8")
        set_cell_margins(cell_lbl)

        # Value cell
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(2)
        p_val.paragraph_format.space_before = Pt(2)
        r_val = p_val.add_run(str(val))
        r_val.font.size = Pt(10)
        set_cell_margins(cell_val)

    # Objective Section
    doc.add_paragraph()
    p_obj_hdr = doc.add_paragraph()
    r_obj_hdr = p_obj_hdr.add_run("OBJETIVO CIENTÍFICO / OPERATIVO:")
    r_obj_hdr.bold = True
    r_obj_hdr.font.size = Pt(10)
    
    p_obj = doc.add_paragraph()
    p_obj.paragraph_format.left_indent = Inches(0.2)
    r_obj = p_obj.add_run(cruise.objective or "No se ha definido un objetivo para esta campaña.")
    r_obj.font.italic = True
    r_obj.font.size = Pt(10)

    # 3. Participants
    doc.add_paragraph()
    p_sec2 = doc.add_paragraph()
    run_sec2 = p_sec2.add_run("2. PARTICIPANTES Y PERSONAL A BORDO")
    run_sec2.bold = True
    run_sec2.font.size = Pt(12)
    run_sec2.font.color.rgb = RGBColor(0x0A, 0x26, 0x47)

    # Split into Crew and Science
    tripulacion = cruise.crew or []
    cientificos = cruise.participants or []

    # 3.1 Tripulación
    p_trip_hdr = doc.add_paragraph()
    r_trip_hdr = p_trip_hdr.add_run("⚓ Tripulación")
    r_trip_hdr.bold = True
    r_trip_hdr.font.size = Pt(11)

    if tripulacion:
        t_trip = doc.add_table(rows=1, cols=5)
        style_table(t_trip)
        hdr_cells = t_trip.rows[0].cells
        headers = ["No.", "Nombre Completo", "Función", "Institución", "Nacionalidad"]
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(cell, "0A2647")
            set_cell_margins(cell, top=120, bottom=120)

        for i, cp in enumerate(tripulacion):
            row = t_trip.add_row()
            cells = row.cells
            
            cells[0].paragraphs[0].add_run(str(i+1)).font.size = Pt(9.5)
            cells[1].paragraphs[0].add_run(cp.personnel.full_name if cp.personnel else "—").font.size = Pt(9.5)
            role_key = cp.role.value if hasattr(cp.role, "value") else str(cp.role)
            role_label = ROLE_LABELS.get(role_key, role_key)
            cells[2].paragraphs[0].add_run(role_label).font.size = Pt(9.5)
            cells[3].paragraphs[0].add_run("CICESE").font.size = Pt(9.5)
            cells[4].paragraphs[0].add_run(cp.personnel.nationality if cp.personnel and cp.personnel.nationality else "—").font.size = Pt(9.5)

            for cell in cells:
                set_cell_margins(cell)
    else:
        doc.add_paragraph("No se han asignado tripulantes a este crucero.").paragraph_format.left_indent = Inches(0.2)

    # 3.2 Científicos
    doc.add_paragraph()
    p_sci_hdr = doc.add_paragraph()
    r_sci_hdr = p_sci_hdr.add_run("🔬 Personal Científico y Técnico")
    r_sci_hdr.bold = True
    r_sci_hdr.font.size = Pt(11)

    if cientificos:
        t_sci = doc.add_table(rows=1, cols=5)
        style_table(t_sci)
        hdr_cells = t_sci.rows[0].cells
        headers = ["No.", "Nombre Completo", "Función / Rol", "Institución", "Nacionalidad"]
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(cell, "1B4F72")
            set_cell_margins(cell, top=120, bottom=120)

        for i, cp in enumerate(cientificos):
            row = t_sci.add_row()
            cells = row.cells
            
            cells[0].paragraphs[0].add_run(str(i+1)).font.size = Pt(9.5)
            r_name = cells[1].paragraphs[0].add_run(cp.participant.full_name if cp.participant else "—")
            r_name.font.size = Pt(9.5)
            if cp.is_principal_investigator:
                r_name.bold = True
                cells[1].paragraphs[0].add_run(" (IP)").bold = True
            role_key = cp.role_in_cruise.value if hasattr(cp.role_in_cruise, "value") else str(cp.role_in_cruise)
            role_label = ROLE_LABELS.get(role_key, role_key)
            cells[2].paragraphs[0].add_run(role_label).font.size = Pt(9.5)
            cells[3].paragraphs[0].add_run(cp.participant.institution if cp.participant else "—").font.size = Pt(9.5)
            cells[4].paragraphs[0].add_run(cp.participant.nationality if cp.participant else "—").font.size = Pt(9.5)

            for cell in cells:
                set_cell_margins(cell)
    else:
        doc.add_paragraph("No se han asignado investigadores ni técnicos a este crucero.").paragraph_format.left_indent = Inches(0.2)

    # 4. Derrotero (Waypoints)
    doc.add_paragraph()
    p_sec3 = doc.add_paragraph()
    run_sec3 = p_sec3.add_run("3. DERROTERO Y ESTACIONES DE MUESTREO")
    run_sec3.bold = True
    run_sec3.font.size = Pt(12)
    run_sec3.font.color.rgb = RGBColor(0x0A, 0x26, 0x47)

    # Insert static map image if available
    map_img = get_static_map_image(cruise.waypoints, cruise.departure_port_ref, cruise.return_port_ref)
    if map_img:
        try:
            p_map = doc.add_paragraph()
            p_map.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_map = p_map.add_run()
            r_map.add_picture(map_img, width=Inches(6.0))
            
            p_map_lbl = doc.add_paragraph()
            p_map_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_map_lbl = p_map_lbl.add_run("Mapa general de la derrota y estaciones de la campaña")
            r_map_lbl.font.size = Pt(8.5)
            r_map_lbl.font.italic = True
            r_map_lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph() # Add space
        except Exception:
            pass

    # Precalcular distancias y tiempos de arribo/salida
    calculated_wps = []
    if cruise.waypoints:
        import math
        from datetime import timedelta
        
        def get_distance_nm(lat1, lon1, lat2, lon2):
            r = 6371.0 # radio en km
            lat1_rad = math.radians(lat1)
            lon1_rad = math.radians(lon1)
            lat2_rad = math.radians(lat2)
            lon2_rad = math.radians(lon2)
            
            dlat = lat2_rad - lat1_rad
            dlon = lon2_rad - lon1_rad
            
            a = math.sin(dlat / 2)**2 + math.cos(lat1_rad) * math.cos(lat2_rad) * math.sin(dlon / 2)**2
            c = 2 * math.asin(math.sqrt(a))
            dist_km = r * c
            return dist_km / 1.852 # mn

        max_speed = cruise.vessel.max_speed_knots if (cruise.vessel and cruise.vessel.max_speed_knots) else 10
        
        # Usar zarpe del modelo (datetime)
        current_dt = None
        if cruise.departure_date:
            d = cruise.departure_date
            if isinstance(d, datetime):
                current_dt = d
            else:
                current_dt = datetime(d.year, d.month, d.day, 8, 0, 0)

        def format_dt(dt):
            if not dt:
                return "—"
            return dt.strftime("%d/%m %H:%M")

        # Construir lista de puntos de trayecto
        trip_points = []
        
        # 1. Puerto de Salida
        dp = cruise.departure_port_ref
        if dp and dp.latitude is not None and dp.longitude is not None:
            trip_points.append({
                "is_port": True,
                "is_departure": True,
                "name": f"Salida de {dp.name}",
                "latitude": dp.latitude,
                "longitude": dp.longitude,
                "speed_knots": None,
                "duration_hours": None,
                "activity": "Salida de puerto"
            })
            
        # 2. Waypoints reales
        for idx, wp in enumerate(cruise.waypoints):
            trip_points.append({
                "is_port": False,
                "is_departure": False,
                "name": wp.name or f"Waypoint {idx+1}",
                "latitude": wp.latitude,
                "longitude": wp.longitude,
                "speed_knots": wp.speed_knots,
                "duration_hours": wp.duration_hours,
                "activity": wp.activity or wp.description or "—"
            })
            
        # 3. Puerto de Regreso
        rp = cruise.return_port_ref
        if rp and rp.latitude is not None and rp.longitude is not None:
            trip_points.append({
                "is_port": True,
                "is_departure": False,
                "name": f"Regreso a {rp.name}",
                "latitude": rp.latitude,
                "longitude": rp.longitude,
                "speed_knots": None,
                "duration_hours": None,
                "activity": "Llegada a puerto"
            })

        cum_dist = 0.0
        for i, pt in enumerate(trip_points):
            arrival_str = "—"
            departure_str = "—"
            
            if i == 0:
                # Primer punto (Inicio del viaje)
                if current_dt:
                    arrival_str = "—"
                    departure_str = format_dt(current_dt)
                    if pt["duration_hours"]:
                        current_dt += timedelta(hours=pt["duration_hours"])
                        departure_str = format_dt(current_dt)
            else:
                prev_pt = trip_points[i - 1]
                dist = get_distance_nm(prev_pt["latitude"], prev_pt["longitude"], pt["latitude"], pt["longitude"])
                cum_dist += dist
                
                if current_dt:
                    speed = pt["speed_knots"] if pt["speed_knots"] else max_speed
                    transit_hours = dist / speed
                    
                    # Llegada = Salida de anterior + tránsito
                    current_dt += timedelta(hours=transit_hours)
                    arrival_str = format_dt(current_dt)
                    
                    # Salida = Llegada + duración de actividad
                    duration = pt["duration_hours"] if pt["duration_hours"] else 0.0
                    current_dt += timedelta(hours=duration)
                    departure_str = format_dt(current_dt)
            
            # Si es el puerto de regreso, no hay hora de salida
            if pt.get("is_port") and not pt.get("is_departure"):
                departure_str = "—"

            calculated_wps.append({
                "name": pt["name"],
                "coords": f"{pt['latitude']:.4f}°, {pt['longitude']:.4f}°",
                "cum_dist_str": f"{cum_dist:.1f} mn",
                "arrival_str": arrival_str,
                "departure_str": departure_str,
                "activity": pt["activity"]
            })

    if calculated_wps:
        t_wp = doc.add_table(rows=1, cols=7)
        style_table(t_wp)
        hdr_cells = t_wp.rows[0].cells
        headers = ["No.", "Estación / Punto", "Coordenadas", "Dist. Acum.", "Arribo (Est.)", "Salida (Est.)", "Actividad"]
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(cell, "2C74B3")
            set_cell_margins(cell, top=120, bottom=120)

        for i, wp in enumerate(calculated_wps):
            row = t_wp.add_row()
            cells = row.cells
            
            cells[0].paragraphs[0].add_run(str(i+1)).font.size = Pt(9)
            cells[1].paragraphs[0].add_run(wp["name"]).font.size = Pt(9)
            cells[2].paragraphs[0].add_run(wp["coords"]).font.size = Pt(9)
            cells[3].paragraphs[0].add_run(wp["cum_dist_str"]).font.size = Pt(9)
            cells[4].paragraphs[0].add_run(wp["arrival_str"]).font.size = Pt(9)
            cells[5].paragraphs[0].add_run(wp["departure_str"]).font.size = Pt(9)
            cells[6].paragraphs[0].add_run(wp["activity"]).font.size = Pt(9)

            for cell in cells:
                set_cell_margins(cell)
    else:
        doc.add_paragraph("No hay puntos o estaciones registradas en la ruta de este crucero.").paragraph_format.left_indent = Inches(0.2)

    # 4. Matriz de Muestreo
    has_samples = any(wp.samples for wp in cruise.waypoints) if cruise.waypoints else False
    if has_samples:
        doc.add_paragraph()
        p_sec4 = doc.add_paragraph()
        r_sec4 = p_sec4.add_run("4. MATRIZ DE MUESTREO CIENTÍFICO POR ESTACIÓN")
        r_sec4.bold = True
        r_sec4.font.size = Pt(12)
        r_sec4.font.color.rgb = RGBColor(0x16, 0xA0, 0x85)

        t_samp = doc.add_table(rows=1, cols=6)
        style_table(t_samp)
        hdr_cells = t_samp.rows[0].cells
        headers = ["Estación", "Ord.", "Variable a Analizar", "Científico Responsable", "Volumen", "Niveles Prof."]
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(cell, "16A085")
            set_cell_margins(cell, top=120, bottom=120)

        for wp in cruise.waypoints:
            if not wp.samples:
                continue
            for j, sample in enumerate(wp.samples):
                row = t_samp.add_row()
                cells = row.cells
                
                cells[0].paragraphs[0].add_run(wp.name or f"Estación" if j == 0 else "").font.size = Pt(9)
                cells[1].paragraphs[0].add_run(str(sample.sampling_order)).font.size = Pt(9)
                r_var = cells[2].paragraphs[0].add_run(sample.variable_name)
                r_var.font.size = Pt(9)
                r_var.bold = True
                cells[3].paragraphs[0].add_run(sample.responsible_name or "—").font.size = Pt(9)
                cells[4].paragraphs[0].add_run(sample.volume_needed or "—").font.size = Pt(9)
                
                depths = []
                if sample.depth_surface: depths.append("Sup")
                if sample.depth_mid_water: depths.append("Med")
                if sample.depth_bottom: depths.append("Fon")
                if sample.depth_custom: depths.append(sample.depth_custom)
                cells[5].paragraphs[0].add_run(", ".join(depths) if depths else "—").font.size = Pt(9)

                for cell in cells:
                    set_cell_margins(cell)

    # 5. Lista de Embarque e Inventario
    if cruise.checklist:
        doc.add_paragraph()
        p_sec5 = doc.add_paragraph()
        r_sec5 = p_sec5.add_run("5. LISTA DE EMBARQUE E INVENTARIO POR INVESTIGADOR")
        r_sec5.bold = True
        r_sec5.font.size = Pt(12)
        r_sec5.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)

        t_chk = doc.add_table(rows=1, cols=4)
        style_table(t_chk)
        hdr_cells = t_chk.rows[0].cells
        headers = ["Investigador Responsable", "Equipo / Material / Reactivo", "Cantidad", "Estado"]
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(cell, "8E44AD")
            set_cell_margins(cell, top=120, bottom=120)

        for item in cruise.checklist:
            row = t_chk.add_row()
            cells = row.cells
            
            cells[0].paragraphs[0].add_run(item.investigator_name).font.size = Pt(9)
            cells[0].paragraphs[0].runs[0].bold = True
            item_desc = item.item_name
            if item.notes:
                item_desc += f" ({item.notes})"
            cells[1].paragraphs[0].add_run(item_desc).font.size = Pt(9)
            cells[2].paragraphs[0].add_run(str(item.quantity)).font.size = Pt(9)
            
            r_status = cells[3].paragraphs[0].add_run("✓ A bordo" if item.is_boarded else "✗ Pendiente")
            r_status.font.size = Pt(9)
            r_status.bold = True
            if item.is_boarded:
                r_status.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            else:
                r_status.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

            for cell in cells:
                set_cell_margins(cell)

    # 6. Plan de Logística y Descargas
    if cruise.discharges:
        doc.add_paragraph()
        p_sec6 = doc.add_paragraph()
        r_sec6 = p_sec6.add_run("6. PLAN DE LOGÍSTICA Y DESCARGA DE MUESTRAS")
        r_sec6.bold = True
        r_sec6.font.size = Pt(12)
        r_sec6.font.color.rgb = RGBColor(0xD3, 0x54, 0x00)

        t_dis = doc.add_table(rows=1, cols=4)
        style_table(t_dis)
        hdr_cells = t_dis.rows[0].cells
        headers = ["Punto / Puerto de Descarga", "Fecha y Hora Estimada", "Contacto en Tierra", "Laboratorio / Notas"]
        for idx, text in enumerate(headers):
            cell = hdr_cells[idx]
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(cell, "D35400")
            set_cell_margins(cell, top=120, bottom=120)

        for d in cruise.discharges:
            row = t_dis.add_row()
            cells = row.cells
            
            cells[0].paragraphs[0].add_run(d.port_name).font.size = Pt(9)
            cells[0].paragraphs[0].runs[0].bold = True
            
            discharge_time = d.discharge_date.strftime("%d/%m/%Y %H:%M") if d.discharge_date else "—"
            cells[1].paragraphs[0].add_run(discharge_time).font.size = Pt(9)
            cells[2].paragraphs[0].add_run(d.responsible_land_person or "—").font.size = Pt(9)
            
            notes_str = d.destination_lab or "—"
            if d.notes:
                notes_str += f" ({d.notes})"
            cells[3].paragraphs[0].add_run(notes_str).font.size = Pt(9)

            for cell in cells:
                set_cell_margins(cell)

    # Notes
    if cruise.notes:
        doc.add_paragraph()
        p_note_hdr = doc.add_paragraph()
        r_note_hdr = p_note_hdr.add_run("NOTAS GENERALES:")
        r_note_hdr.bold = True
        r_note_hdr.font.size = Pt(10)
        
        p_note = doc.add_paragraph()
        p_note.paragraph_format.left_indent = Inches(0.2)
        r_note = p_note.add_run(cruise.notes)
        r_note.font.size = Pt(9.5)

    # Save to stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
